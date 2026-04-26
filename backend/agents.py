import json
import asyncio
import os
from google.genai import Client as GeminiClient
from google.genai.errors import ClientError
from google.genai import types
from groq import AsyncGroq
from docx import Document
from docx.shared import Inches, Pt
import base64
from anthropic import AsyncAnthropic

from config import (
    GEMINI_API_KEYS, TEXT_MODEL, EMBEDDING_MODEL, IMAGE_MODEL,
    GROQ_API_KEYS, GROQ_MODEL,
    ANTHROPIC_API_KEYS, CLAUDE_MODEL
)
from prompts import (
    SECTION_EXTRACTION_PROMPT,
    FULL_SUMMARY_PROMPT,
    SECTION_SUMMARY_PROMPT,
    RESEARCH_DISCOVERY_PROMPT,
    WORKFLOW_GENERATION_PROMPT
)


async def _gemini_call_with_retry(api_keys, model, contents, max_retries=5, response_mime_type=None):
    """Call Gemini with automatic retry and key rotation on rate-limit (429) errors."""
    keys_list = api_keys if isinstance(api_keys, list) else [api_keys]
    max_attempts = max(max_retries, len(keys_list) * 2) # ensure we try all keys
    for attempt in range(max_attempts):
        current_key = keys_list[attempt % len(keys_list)]
        client = GeminiClient(api_key=current_key)
        try:
            config = None
            if response_mime_type:
                config = types.GenerateContentConfig(response_mime_type=response_mime_type)
            
            response = await client.aio.models.generate_content(
                model=model, contents=contents, config=config
            )
            return response
        except Exception as e:
            # Check for 429 (Rate Limit / Resource Exhausted)
            is_rate_limit = False
            if hasattr(e, 'code') and e.code == 429:
                is_rate_limit = True
            elif "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                is_rate_limit = True
                
            if is_rate_limit and attempt < max_attempts - 1:
                # If we have multiple keys and haven't tried all of them yet in this cycle
                if len(keys_list) > 1 and (attempt + 1) % len(keys_list) != 0:
                    print(f"[Gemini Rate limit] Switching to next API key (attempt {attempt + 1}/{max_attempts})")
                    # Don't sleep, just immediately loop and try the next key
                    continue
                else:
                    wait = 15 * ((attempt // len(keys_list)) + 1)
                    print(f"[Gemini Rate limit] All keys exhausted/rate-limited. Sleeping {wait}s... (attempt {attempt + 1}/{max_attempts})")
                    await asyncio.sleep(wait)
            else:
                raise


async def _groq_call_with_retry(api_keys, model, messages, max_retries=3, temperature=0.7, max_tokens=4096, response_format=None):
    """Call Groq with automatic retry and key rotation on rate-limit (429) errors."""
    keys_list = api_keys if isinstance(api_keys, list) else [api_keys]
    max_attempts = max(max_retries, len(keys_list) * 2)
    for attempt in range(max_attempts):
        current_key = keys_list[attempt % len(keys_list)]
        client = AsyncGroq(api_key=current_key)
        try:
            kwargs = {
                "model": model,
                "messages": messages,
                "temperature": temperature,
                "max_tokens": max_tokens,
            }
            if response_format:
                kwargs["response_format"] = response_format
                
            response = await client.chat.completions.create(**kwargs)
            return response
        except Exception as e:
            is_rate_limit = False
            if "429" in str(e) or "rate_limit_exceeded" in str(e).lower():
                is_rate_limit = True
            
            if is_rate_limit and attempt < max_attempts - 1:
                if len(keys_list) > 1 and (attempt + 1) % len(keys_list) != 0:
                    print(f"[Groq Rate limit] Switching to next API key (attempt {attempt + 1}/{max_attempts})")
                    continue
                else:
                    wait = 5 * (2 ** (attempt // len(keys_list)))
                    print(f"[Groq Rate limit] All keys exhausted/rate-limited. Sleeping {wait}s... (attempt {attempt + 1}/{max_attempts})")
                    await asyncio.sleep(wait)
            else:
                raise


async def _claude_call_with_retry(api_keys, model, messages, max_retries=3, temperature=0.6, max_tokens=8192):
    """Call Claude with automatic retry and key rotation on rate-limit (429) errors."""
    keys_list = api_keys if isinstance(api_keys, list) else [api_keys]
    max_attempts = max(max_retries, len(keys_list) * 2)
    for attempt in range(max_attempts):
        current_key = keys_list[attempt % len(keys_list)]
        client = AsyncAnthropic(api_key=current_key)
        try:
            response = await client.messages.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return response
        except Exception as e:
            is_rate_limit = False
            err_str = str(e).lower()
            if "429" in err_str or "rate limit" in err_str or "overloaded" in err_str:
                is_rate_limit = True

            if is_rate_limit and attempt < max_attempts - 1:
                if len(keys_list) > 1 and (attempt + 1) % len(keys_list) != 0:
                    print(f"[Claude Rate limit] Switching to next API key (attempt {attempt + 1}/{max_attempts})")
                    continue
                else:
                    wait = 5 * (2 ** (attempt // len(keys_list)))
                    print(f"[Claude Rate limit] All keys exhausted/rate-limited. Sleeping {wait}s... (attempt {attempt + 1}/{max_attempts})")
                    await asyncio.sleep(wait)
            else:
                raise

class SummarizationAgent:
    """Agent 1 — Powered by Google Gemini.
    Handles heading extraction, full-paper summaries, and section-level summaries.
    """

    def __init__(self, api_keys: list = GEMINI_API_KEYS):
        self.api_keys = api_keys

    async def extract_headings(self, full_text: str) -> list[str]:
        """Ask the LLM to identify section headings from the raw paper text."""
        prompt = SECTION_EXTRACTION_PROMPT.format(text=full_text[:15000])
        response = await _gemini_call_with_retry(
            self.api_keys, 
            TEXT_MODEL, 
            contents=prompt,
            response_mime_type="application/json"
        )
        try:
            raw = response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
            return json.loads(raw)
        except (json.JSONDecodeError, IndexError):
            return ["Full Paper"]

    async def summarize_full(self, context: str) -> str:
        """Generate an overall summary of the entire paper."""
        prompt = FULL_SUMMARY_PROMPT.format(context=context)
        response = await _gemini_call_with_retry(self.api_keys, TEXT_MODEL, contents=prompt)
        return response.text

    async def summarize_section(self, heading: str, content: str) -> str:
        """Summarize a specific section of the paper."""
        prompt = SECTION_SUMMARY_PROMPT.format(heading=heading, content=content)
        response = await _gemini_call_with_retry(self.api_keys, TEXT_MODEL, contents=prompt)
        return response.text


class ResearchDiscoveryAgent:
    """Agent 2 — Research Discovery.
    Generates future research ideas, alternative models, and related papers using Gemini's native JSON mode.
    """

    def __init__(self, api_keys: list = GEMINI_API_KEYS):
        self.api_keys = api_keys

    async def discover(self, context: str) -> dict:
        """Analyze the paper and return structured research discovery insights."""
        prompt = RESEARCH_DISCOVERY_PROMPT.format(context=context)
        try:
            response = await _gemini_call_with_retry(
                api_keys=self.api_keys,
                model=TEXT_MODEL,
                contents=prompt,
                response_mime_type="application/json"
            )
            raw = response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
            return json.loads(raw)
        except Exception as e:
            return {
                "future_ideas": [{"title": "Parse error", "description": str(e)[:500], "impact": "N/A"}],
                "alternative_models": [],
                "related_papers": [],
            }


class WorkflowGeneratorAgent:
    """Agent 3 — Triggered by specific idea selection.
    Uses Groq for detailed project workflow text generation.
    """

    def __init__(self, gemini_keys: list = GEMINI_API_KEYS, groq_keys: list = GROQ_API_KEYS, anthropic_keys: list = ANTHROPIC_API_KEYS):
        self.groq_keys = groq_keys

    async def generate_workflow_text(self, selected_item: str, context: str) -> str:
        """Generate a structured project workflow as text using Groq."""
        prompt_text = WORKFLOW_GENERATION_PROMPT.format(
            selected_item=selected_item,
            context=context[:10000] # Reduced context to avoid TPD limits
        )

        messages = [
            {"role": "system", "content": "You are a senior research project architect. Generate a highly detailed and structured project workflow."},
            {"role": "user", "content": prompt_text}
        ]

        response = await _groq_call_with_retry(
            api_keys=self.groq_keys,
            model=GROQ_MODEL,
            messages=messages,
            temperature=0.6
        )
        return response.choices[0].message.content.strip()

    def create_word_doc(self, workflow_text: str, output_path: str):
        """Convert the Markdown workflow text into a formatted .docx file."""
        doc = Document()
        
        lines = workflow_text.split('\n')
        title = "Research Project Workflow"
        # Try to find title in text
        for line in lines:
            if "Project Title" in line or "PROJECT TITLE" in line:
                if ":" in line:
                    title = line.split(":", 1)[1].strip().strip('*').strip()
                    break
        
        doc.add_heading(title, 0)
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif (line.startswith('- ') or line.startswith('* ')) and len(line) > 2:
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and (line[1] == '.' or (line[2] == '.' if len(line)>2 else False)):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)
        
        doc.save(output_path)
